#!/usr/bin/env python3
"""extract.py - content extraction helpers for the document database.

docx/pdf -> markdown, xlsx -> one csv per sheet.

Used by ingest.py, but also runnable standalone:

    python3 scripts/extract.py path/to/file.docx [output_dir]
"""

import csv
import io
import sys
from pathlib import Path

EXTRACTABLE_TYPES = {".docx", ".pdf", ".xlsx"}


def extract_docx(path: Path) -> str:
    """Extract a .docx file to markdown (headings, paragraphs, tables)."""
    import docx

    doc = docx.Document(str(path))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if style.startswith("heading"):
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 2
            lines.append("#" * min(level, 6) + " " + text)
        elif style == "title":
            lines.append("# " + text)
        else:
            lines.append(text)
        lines.append("")
    for table in doc.tables:
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells]
                for row in table.rows]
        if not rows:
            continue
        lines.append("| " + " | ".join(rows[0]) + " |")
        lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def extract_pdf(path: Path) -> str:
    """Extract a .pdf file's text to markdown (one section per page)."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        parts.append(f"## Page {i}\n\n{text}\n")
    return "\n".join(parts).strip() + "\n"


def extract_xlsx(path: Path) -> list[tuple[str, str]]:
    """Extract a .xlsx file. Returns [(sheet_name, csv_text), ...]."""
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    sheets = []
    for ws in wb.worksheets:
        buf = io.StringIO()
        writer = csv.writer(buf)
        for row in ws.iter_rows(values_only=True):
            writer.writerow(["" if v is None else v for v in row])
        sheets.append((ws.title, buf.getvalue()))
    wb.close()
    return sheets


def safe_sheet_name(name: str) -> str:
    """Make a sheet name filesystem-safe for use in csv filenames."""
    return "".join(c if c.isalnum() or c in "-_ " else "_" for c in name).strip()


def extract_to_dir(src: Path, out_dir: Path) -> list[Path]:
    """Extract src into out_dir. Returns the list of files written.

    Naming matches the repo convention:
      report.docx -> report.docx.md
      report.pdf  -> report.pdf.md
      book.xlsx   -> book.xlsx__{sheet}.csv (one per sheet)
    """
    out_dir.mkdir(parents=True, exist_ok=True)
    suffix = src.suffix.lower()
    written = []
    if suffix == ".docx":
        out = out_dir / (src.name + ".md")
        out.write_text(extract_docx(src), encoding="utf-8")
        written.append(out)
    elif suffix == ".pdf":
        out = out_dir / (src.name + ".md")
        out.write_text(extract_pdf(src), encoding="utf-8")
        written.append(out)
    elif suffix == ".xlsx":
        for sheet_name, csv_text in extract_xlsx(src):
            out = out_dir / f"{src.name}__{safe_sheet_name(sheet_name)}.csv"
            out.write_text(csv_text, encoding="utf-8")
            written.append(out)
    return written


def main() -> None:
    if len(sys.argv) < 2:
        sys.exit(__doc__)
    src = Path(sys.argv[1])
    out_dir = Path(sys.argv[2]) if len(sys.argv) > 2 else src.parent
    if src.suffix.lower() not in EXTRACTABLE_TYPES:
        sys.exit(f"unsupported file type: {src.suffix} "
                 f"(supported: {', '.join(sorted(EXTRACTABLE_TYPES))})")
    for path in extract_to_dir(src, out_dir):
        print(path)


if __name__ == "__main__":
    main()
