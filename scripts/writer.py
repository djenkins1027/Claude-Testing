import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas as canvas_module


def _add_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def build_docx(text: str, footer_stamp: str | None = None) -> bytes:
    doc = Document()
    for line in text.splitlines() or [""]:
        if line.startswith("## "):
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("# "):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    if footer_stamp:
        footer_paragraph = doc.sections[0].footer.paragraphs[0]
        footer_paragraph.add_run(f"{footer_stamp}  Page ")
        _add_field(footer_paragraph, "PAGE")
        footer_paragraph.add_run(" of ")
        _add_field(footer_paragraph, "NUMPAGES")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_xlsx(rows: list[list], footer_stamp: str | None = None) -> bytes:
    wb = Workbook()
    ws = wb.active
    for row in rows:
        ws.append(row)

    if footer_stamp:
        ws.oddFooter.left.text = footer_stamp
        ws.oddFooter.right.text = "Page &P of &N"

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


class _NumberedCanvas(canvas_module.Canvas):
    """Buffers pages so the footer can show 'Page X of Y' (total pages isn't known until the end)."""

    def __init__(self, *args, footer_stamp=None, **kwargs):
        super().__init__(*args, **kwargs)
        self.footer_stamp = footer_stamp
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        total_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self._draw_footer(total_pages)
            super().showPage()
        super().save()

    def _draw_footer(self, total_pages):
        if not self.footer_stamp:
            return
        width, _ = letter
        self.setFont("Helvetica", 9)
        self.drawString(72, 30, self.footer_stamp)
        self.drawRightString(width - 72, 30, f"Page {self._pageNumber} of {total_pages}")


def build_pdf(text: str, footer_stamp: str | None = None) -> bytes:
    buffer = io.BytesIO()
    pdf = _NumberedCanvas(buffer, pagesize=letter, footer_stamp=footer_stamp)
    width, height = letter
    y = height - 72
    for line in text.splitlines() or [""]:
        pdf.drawString(72, y, line)
        y -= 14
        if y < 72:
            pdf.showPage()
            y = height - 72
    pdf.showPage()  # flush the final (not-yet-wrapped) page into the buffered states
    pdf.save()
    return buffer.getvalue()
